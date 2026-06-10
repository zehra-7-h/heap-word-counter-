#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Raporu Word formatına dönüştür ve İndirilenler'e kaydet
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Rapor dosyasını oku
rapor_path = r"c:\Users\Administrator\OneDrive\Belgeler\Desktop\projelerim\LAB PROJE SUNUMLAR\LABPROJESUNUM6\PROJE_RAPORU_RESMI.txt"
downloads_path = r"d:\Documents\Downloads\PROJE_RAPORU_RESMI.docx"

# Dosyayı oku
with open(rapor_path, 'r', encoding='utf-8') as f:
    content = f.read()

# Word dokümanı oluştur
doc = Document()

# Başlık ekle
title = doc.add_paragraph()
title_run = title.add_run("T.C\nKOCAELİ SAĞLIK VE TEKNOLOJİ ÜNİVERSİTESİ\nMÜHENDİSLİK FAKÜLTESİ\nBİLGİSAYAR MÜHENDİSLİĞİ BÖLÜMÜ")
title_run.font.size = Pt(12)
title_run.font.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Boş satır
doc.add_paragraph()

# Proje başlığı
proje_baslik = doc.add_paragraph()
proje_run = proje_baslik.add_run("PROJE BAŞLIGI:\nHEAP YAPISINA DAYALI KELİME SAYICI PROGRAM")
proje_run.font.size = Pt(14)
proje_run.font.bold = True
proje_baslik.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Boş satır
doc.add_paragraph()

# İçeriği bölümlere ayırarak ekle
lines = content.split('\n')
for line in lines:
    if line.strip():
        # Başlık satırları
        if line.startswith(('1 ', '2 ', '3 ', '4 ', '5 ', '6 ')) and len(line) < 50:
            p = doc.add_paragraph(line.strip(), style='Heading 1')
            p_format = p.paragraph_format
            p_format.space_before = Pt(12)
            p_format.space_after = Pt(6)
        
        # Alt başlıklar
        elif line.startswith(('1.1', '1.2', '2.1', '2.2', '2.3', '3.1', '3.2', '3.3', '3.4', '4.1', '4.2', '5.1', '5.2')):
            p = doc.add_paragraph(line.strip(), style='Heading 2')
            p_format = p.paragraph_format
            p_format.space_before = Pt(10)
            p_format.space_after = Pt(4)
        
        # Bullet points
        elif line.strip().startswith(('├─', '├', '└', '•', '✓', '-', '✅')):
            text = line.strip()
            # Karakterleri temizle
            text = text.replace('├─', '').replace('├', '').replace('└', '').replace('•', '').replace('✓', '').replace('✅', '').strip()
            if text:
                p = doc.add_paragraph(text, style='List Bullet')
                p_format = p.paragraph_format
                p_format.left_indent = Inches(0.5)
        
        else:
            # Normal paragraf
            p = doc.add_paragraph(line.strip())
            p_format = p.paragraph_format
            p_format.space_after = Pt(6)

# Dosyayı kaydet
doc.save(downloads_path)

print(f"✅ Word dosyası başarıyla oluşturuldu!")
print(f"📍 Konum: {downloads_path}")
print(f"📊 Dosya Boyutu: {os.path.getsize(downloads_path) / 1024:.2f} KB")
